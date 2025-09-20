#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import json
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE


def clear_body(doc: Document) -> None:
    for p in reversed(doc.paragraphs):
        p_element = p._element
        p_element.getparent().remove(p_element)
    for t in reversed(doc.tables):
        t_element = t._element
        t_element.getparent().remove(t_element)


def _apply_paragraph_alignment(paragraph, align: str | None) -> None:
    if not align:
        return
    amap = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }
    paragraph.paragraph_format.alignment = amap.get(str(align).lower())


def _apply_paragraph_style(paragraph, root_doc, tfp: dict | None):
    if not isinstance(tfp, dict):
        return
    style_name = tfp.get("paragraphStyleName")
    style_id = tfp.get("paragraphStyleId")
    # Prefer name, fallback to ID
    style_key = style_name or style_id
    if isinstance(style_key, str) and style_key.strip():
        try:
            paragraph.style = root_doc.styles[style_key]
            return
        except KeyError:
            # Fallback for case-insensitivity or name/id mismatch
            s_lower = style_key.lower()
            for s in root_doc.styles:
                if (s.name or "").lower() == s_lower or s.style_id.lower() == s_lower:
                    paragraph.style = s
                    return


def _apply_paragraph_numbering(paragraph, numPr: dict | None) -> None:
    if not isinstance(numPr, dict):
        return
    numId = numPr.get("numId")
    ilvl = numPr.get("ilvl")
    if numId is None or ilvl is None:
        return
    pPr_elm = paragraph._p.get_or_add_pPr()
    numPr_elm = pPr_elm.get_or_add_numPr()
    numId_elm = numPr_elm.get_or_add_numId()
    numId_elm.val = int(numId)
    ilvl_elm = numPr_elm.get_or_add_ilvl()
    ilvl_elm.val = int(ilvl)


def _apply_run_from_tfp(run, tfp: dict | None) -> None:
    if not isinstance(tfp, dict):
        return
    rPr = deepcopy(tfp.get("rPr") or {})
    font = run.font
    # booleans
    if rPr.get("bold"):
        font.bold = True
    if rPr.get("italic"):
        font.italic = True
    u = rPr.get("underline")
    if isinstance(u, dict) and u.get("val") and u.get("val") != "none":
        font.underline = WD_UNDERLINE.SINGLE
    if rPr.get("strike"):
        font.strike = True
    # color
    color = rPr.get("color")
    if isinstance(color, str) and color:
        try:
            font.color.rgb = RGBColor.from_string(color)
        except Exception:
            pass
    # size
    szhp = rPr.get("sizeHalfPoints")
    if isinstance(szhp, int):
        font.size = Pt(szhp / 2)
    else:
        # fallback: fontSizePt as float
        fspt = tfp.get("fontSizePt")
        if isinstance(fspt, (int, float)):
            font.size = Pt(float(fspt))
    # fonts
    font_name = tfp.get("fontName")
    if isinstance(font_name, str) and font_name.strip():
        font.name = font_name
    else:
        rfonts = (rPr.get("rFonts") or {})
        name = rfonts.get("ascii") or rfonts.get("hAnsi")
        if isinstance(name, str) and name.strip():
            font.name = name


def _resolve_run_tfp(text_props_registry: dict, run_obj: dict) -> dict | None:
    tfp = run_obj.get("text_full_properties")
    if isinstance(tfp, dict):
        return tfp
    ref = run_obj.get("textPropsRef")
    if isinstance(ref, str):
        return deepcopy((text_props_registry or {}).get(ref) or {})
    return None


def reconstruct_paragraph(parent, p_obj: dict, text_props_registry: dict | None, root_doc: Document) -> None:
    # parent can be Document or _Cell
    p = parent.add_paragraph()
    # Alignment and numbering from p.p
    pmeta = p_obj.get("p") or {}
    first_run_tfp = _resolve_run_tfp(text_props_registry, (p_obj.get("content", []) or [{}])[0])
    _apply_paragraph_style(p, root_doc, first_run_tfp)
    _apply_paragraph_alignment(p, pmeta.get("jc") or (first_run_tfp or {}).get("alignment"))
    _apply_paragraph_numbering(p, (pmeta.get("numPr") or {}))

    # Default python-docx adds an empty run; remove it
    if p.runs:
        p._element.remove(p.runs[0]._element)

    for item in p_obj.get("content", []) or []:
        if not isinstance(item, dict):
            continue
        if item.get("type") == "run":
            if "text" in item:
                r = p.add_run(item.get("text", ""))
                _apply_run_from_tfp(r, _resolve_run_tfp(text_props_registry, item))
            elif "chunks" in item:
                chunks = item.get("chunks", []) or []
                for ch in chunks:
                    if ch.get("type") == "text":
                        r = p.add_run(ch.get("text", ""))
                        _apply_run_from_tfp(r, _resolve_run_tfp(text_props_registry, item))
                    # Other chunk types (tab, break, etc.) can be added here if needed
        elif item.get("type") == "hyperlink":
            # Simplify: render hyperlink runs as plain runs with properties
            for ritem in item.get("runs", []) or []:
                if isinstance(ritem, dict):
                    if "text" in ritem:
                        r = p.add_run(ritem.get("text", ""))
                        _apply_run_from_tfp(r, _resolve_run_tfp(text_props_registry, ritem))
                    elif "chunks" in ritem:
                        for ch in ritem.get("chunks", []) or []:
                            if ch.get("type") == "text":
                                r = p.add_run(ch.get("text", ""))
                                _apply_run_from_tfp(r, _resolve_run_tfp(text_props_registry, ritem))


def reconstruct_table(parent, t_obj: dict, text_props_registry: dict | None, root_doc: Document) -> None:
    rows_data = t_obj.get("rows", []) or []
    if not rows_data:
        return
    # Determine cols from first row
    num_rows = len(rows_data)
    num_cols = 0
    if num_rows > 0:
        for cell in rows_data[0].get("cells", []) or []:
            num_cols += (cell.get("tcPr", {}).get("gridSpan") or 1)
    if num_rows == 0 or num_cols == 0:
        return
    table = parent.add_table(rows=num_rows, cols=num_cols)
    # Track merged cells to avoid writing to them
    merged_cells = set()
    vMerge_starts = {}

    for i, row in enumerate(rows_data):
        row_cells = table.rows[i].cells
        cell_cursor = 0
        for j, cell in enumerate(row.get("cells", []) or []):
            while (i, cell_cursor) in merged_cells:
                cell_cursor += 1
            if cell_cursor >= num_cols:
                continue

            dcell = row_cells[cell_cursor]
            tcPr = cell.get("tcPr", {})
            gridSpan = tcPr.get("gridSpan")
            if isinstance(gridSpan, int) and gridSpan > 1:
                end_cell_idx = cell_cursor + gridSpan - 1
                if end_cell_idx < num_cols:
                    dcell.merge(row_cells[end_cell_idx])
                    for k in range(cell_cursor + 1, end_cell_idx + 1):
                        merged_cells.add((i,k))
            
            vMerge = tcPr.get("vMerge")
            if vMerge == "restart":
                # Start of a vertical merge. Other cells will refer back to this one.
                vMerge_starts[cell_cursor] = dcell
            elif vMerge == "continue":
                # This cell should be part of a merge from above.
                if cell_cursor in vMerge_starts:
                    vMerge_starts[cell_cursor].merge(dcell)

            # Remove default paragraph
            if dcell.paragraphs:
                p0 = dcell.paragraphs[0]
                p0._element.getparent().remove(p0._element)
            for elem in cell.get("content", []) or []:
                if isinstance(elem, dict) and elem.get("type") == "paragraph":
                    reconstruct_paragraph(dcell, elem, text_props_registry, root_doc)
                elif isinstance(elem, dict) and elem.get("type") == "table":
                    # Nested tables
                    reconstruct_table(dcell, elem, text_props_registry, root_doc)
            
            cell_cursor += (gridSpan or 1)


def reconstruct_body(doc: Document, body: list, text_props_registry: dict | None) -> None:
    for blk in body or []:
        if not isinstance(blk, dict):
            continue
        if blk.get("type") == "paragraph":
            reconstruct_paragraph(doc, blk, text_props_registry, doc)
        elif blk.get("type") == "table":
            reconstruct_table(doc, blk, text_props_registry, doc)


def main():
    if len(sys.argv) != 3:
        print("Usage: python json2docx_structured.py <input.json> <output.docx>", file=sys.stderr)
        sys.exit(2)
    json_path = sys.argv[1]
    out_docx = sys.argv[2]
    if not os.path.isfile(json_path):
        print(f"Input JSON not found: {json_path}", file=sys.stderr)
        sys.exit(2)
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    source_docx = data.get("source_file")
    if not (isinstance(source_docx, str) and os.path.isfile(source_docx)):
        print("Error: source_file missing or not found in JSON; cannot preserve numbering/styles.", file=sys.stderr)
        sys.exit(2)
    try:
        doc = Document(source_docx)
    except Exception as e:
        print(f"Error opening source document: {e}", file=sys.stderr)
        sys.exit(2)
    clear_body(doc)
    text_props_registry = (data.get("textProperties") or {})
    reconstruct_body(doc, data.get("body") or [], text_props_registry)
    try:
        doc.save(out_docx)
        print(f"Wrote {out_docx}")
    except Exception as e:
        print(f"Error saving docx: {e}", file=sys.stderr)
        sys.exit(2)


if __name__ == "__main__":
    main()


